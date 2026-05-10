"""
Generate a Word bulletin (.docx) from an OrderOfWorship by replacing
placeholders in a template document.

The template IS the theme — each church has their own .docx template
with their formatting, spacing, fonts, and layout. This code just
swaps the placeholders with real content and replaces the theme image.
"""

import copy
import io
import re
from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .models import OrderOfWorship
from .paths import RESOURCES_DIR, OUTPUT_DIR
from . import calendar_data

# Olive accent color for calendar block
CALENDAR_OLIVE = RGBColor(0x6B, 0x7A, 0x3D)
CALENDAR_FONT = "Georgia"

# Template path — this is the "theme"
TEMPLATE_PATH = RESOURCES_DIR / "bulletin" / "Template - Bulletin.docx"


def _ordinal_date(d: date) -> str:
    """Format a date like 'April 5th, 2026' with ordinal suffix."""
    day = d.day
    if 11 <= day <= 13:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    return f"{d.strftime('%B')} {day}{suffix}, {d.year}"


def _liturgical_prayer_label(order: OrderOfWorship) -> str:
    """Build the liturgical prayer label from the selected prayer."""
    if not order.liturgicalPrayer:
        return 'THE LORD\u2019S PRAYER'

    num = order.liturgicalPrayer.number
    title = order.liturgicalPrayer.title

    # Shorten the common Lord's Prayer variants
    if num in ('894', '895', '896'):
        return "THE LORD\u2019S PRAYER"

    return title.upper() if title else 'THE LORD\u2019S PRAYER'


def _replace_in_runs(paragraph, placeholder: str, replacement: str):
    """
    Replace a placeholder across paragraph runs, preserving formatting.

    Word splits text across multiple runs unpredictably. This finds the
    placeholder across run boundaries and replaces it while keeping the
    formatting of the first run that contains part of the placeholder.
    """
    # First, try simple per-run replacement
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)
            return True

    # If not found in a single run, search across runs
    full_text = ''.join(run.text for run in paragraph.runs)
    if placeholder not in full_text:
        return False

    # Find where the placeholder starts and ends across runs
    idx = full_text.index(placeholder)
    end_idx = idx + len(placeholder)

    # Rebuild runs: before placeholder, replacement, after placeholder
    char_pos = 0
    new_runs_text = []
    replacement_inserted = False

    for run in paragraph.runs:
        run_start = char_pos
        run_end = char_pos + len(run.text)

        if run_end <= idx:
            # Entirely before placeholder — keep as is
            new_runs_text.append(None)
        elif run_start >= end_idx:
            # Entirely after placeholder — keep as is
            new_runs_text.append(None)
        else:
            # This run overlaps with the placeholder
            before = run.text[:max(0, idx - run_start)]
            after = run.text[max(0, end_idx - run_start):]

            if not replacement_inserted:
                new_runs_text.append(before + replacement + after)
                replacement_inserted = True
            else:
                # Subsequent runs that were part of the placeholder
                new_runs_text.append(after if after else '')

        char_pos = run_end

    # Apply the new text to runs
    for run, new_text in zip(paragraph.runs, new_runs_text):
        if new_text is not None:
            run.text = new_text

    return True


def _replace_theme_image(doc: Document, image_path: Path):
    """Replace the first inline image in the document with the theme image."""
    if not image_path or not image_path.exists():
        return

    # Convert unsupported formats to PNG
    supported = {'.bmp', '.gif', '.jpg', '.jpeg', '.png', '.tiff', '.tif', '.wmf'}
    if image_path.suffix.lower() not in supported:
        try:
            from PIL import Image
            png_path = image_path.with_suffix('.png')
            if not png_path.exists():
                img = Image.open(image_path)
                img.save(png_path, 'PNG')
            image_path = png_path
        except Exception:
            return

    # Find the first inline shape (the theme image placeholder)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                # Found an inline image — get its relationship ID
                blips = run._element.findall(
                    './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
                )
                if blips:
                    blip = blips[0]
                    ns = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'
                    rId = blip.get(f'{ns}embed')
                    if rId:
                        # Replace the image data in the relationship
                        part = doc.part
                        image_part = part.related_parts[rId]
                        with open(image_path, 'rb') as f:
                            image_part._blob = f.read()
                        return


def _format_event_date(d: date) -> tuple[str, str, str, str]:
    """Return (DAYNAME, 'Month Day', suffix, '') split for superscript rendering."""
    day = d.day
    if 11 <= day <= 13:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    return d.strftime("%A").upper(), f"{d.strftime('%B')} {day}", suffix, ""


def _add_olive_left_border(paragraph):
    """Add a thick olive left border to a paragraph for visual accent."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')  # 3pt thick
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '6B7A3D')
    p_bdr.append(left)
    p_pr.append(p_bdr)


def _add_olive_bottom_border(paragraph):
    """Add a thick olive bottom border (horizontal divider)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # 1.5pt thick (lighter than vertical bar)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '6B7A3D')
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _read_anchor_font(paragraph) -> tuple[str, float]:
    """Read font name and size from the {{CALENDAR_BLOCK}} placeholder."""
    font_name = CALENDAR_FONT
    font_size = 11.0
    for run in paragraph.runs:
        if run.font.name:
            font_name = run.font.name
        if run.font.size:
            font_size = run.font.size.pt
            break
    return font_name, font_size


def _render_calendar_block(doc, anchor_paragraph, service_date_str: str):
    """
    Replace the {{CALENDAR_BLOCK}} paragraph with a styled calendar.

    Font name/size is inherited from the placeholder paragraph.
    Day headers are +2pt over the base.
    Olive accent for headers and times.
    """
    cal = calendar_data.get_calendar_for_service(service_date_str)
    events = cal.get("events", [])

    # Read font from anchor paragraph
    base_font, base_size = _read_anchor_font(anchor_paragraph)
    header_size = base_size + 2

    # Group events by date
    by_date: dict[str, list] = {}
    for e in events:
        by_date.setdefault(e["date"], []).append(e)

    # Get the parent and index for inserting after the anchor
    p_element = anchor_paragraph._p
    parent = p_element.getparent()
    anchor_index = list(parent).index(p_element)

    new_paragraphs = []

    # Track approximate space usage for overflow warning
    # Page 4 has ~7" of vertical space; each event is ~0.25", date header ~0.4"
    estimated_height = 0.0

    for i, (date_str, day_events) in enumerate(sorted(by_date.items())):
        d = date.fromisoformat(date_str)
        day_name, date_label, ordinal_suffix, _ = _format_event_date(d)

        # Spacer paragraph between weeks (so left border doesn't bleed into space-before)
        if i > 0:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(0)
            sp_run = spacer.add_run("")
            sp_run.font.size = Pt(8)
            new_paragraphs.append(spacer)

        # Date header — uses left border accent instead of bottom line
        header_p = doc.add_paragraph()
        header_p.paragraph_format.space_before = Pt(0)
        header_p.paragraph_format.space_after = Pt(4)
        header_p.paragraph_format.left_indent = Inches(0.15)

        r1 = header_p.add_run(day_name)
        r1.font.name = base_font
        r1.font.size = Pt(header_size)
        r1.font.bold = True
        r1.font.color.rgb = CALENDAR_OLIVE

        r2 = header_p.add_run(f"  ·  {date_label}")
        r2.font.name = base_font
        r2.font.size = Pt(header_size)
        r2.font.bold = True
        r2.font.color.rgb = CALENDAR_OLIVE

        r3 = header_p.add_run(ordinal_suffix)
        r3.font.name = base_font
        r3.font.size = Pt(header_size)
        r3.font.bold = True
        r3.font.color.rgb = CALENDAR_OLIVE
        r3.font.superscript = True

        _add_olive_left_border(header_p)
        new_paragraphs.append(header_p)
        estimated_height += 0.45

        # Events for this date
        # Tab stops: time | title | location aligned at fixed columns
        for event in day_events:
            ev_p = doc.add_paragraph()
            ev_p.paragraph_format.space_before = Pt(2)
            ev_p.paragraph_format.space_after = Pt(0)
            ev_p.paragraph_format.left_indent = Inches(0.4)

            # Time at 0.4", title at 1.5", location at 3.6"
            tab_stops = ev_p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(1.5))
            tab_stops.add_tab_stop(Inches(3.6))

            time_run = ev_p.add_run(event.get("time", ""))
            time_run.font.name = base_font
            time_run.font.size = Pt(base_size)
            time_run.font.bold = True
            time_run.font.color.rgb = CALENDAR_OLIVE

            title_run = ev_p.add_run(f"\t{event.get('title', '')}")
            title_run.font.name = base_font
            title_run.font.size = Pt(base_size)

            location = event.get("location", "")
            if location:
                loc_run = ev_p.add_run(f"\t{location}")
                loc_run.font.name = base_font
                loc_run.font.size = Pt(max(base_size - 1, 8))
                loc_run.font.italic = True
                loc_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

            new_paragraphs.append(ev_p)
            estimated_height += 0.25

    # Add the optional note if enabled
    note = calendar_data.get_note(service_date_str)
    if note.get("enabled") and note.get("text", "").strip():
        # Ornate olive divider: ──────── ❦ ────────
        divider_p = doc.add_paragraph()
        divider_p.paragraph_format.space_before = Pt(18)
        divider_p.paragraph_format.space_after = Pt(0)
        divider_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        ornament_text = "────────  ❦  ────────"
        div_run = divider_p.add_run(ornament_text)
        div_run.font.name = base_font
        div_run.font.size = Pt(base_size + 1)
        div_run.font.color.rgb = CALENDAR_OLIVE

        new_paragraphs.append(divider_p)

        note_p = doc.add_paragraph()
        note_p.paragraph_format.space_before = Pt(10)
        note_p.paragraph_format.space_after = Pt(0)

        # Apply justification
        align = note.get("align", "left")
        if align == "center":
            note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            note_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            note_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        label_run = note_p.add_run("Note: ")
        label_run.font.name = base_font
        label_run.font.size = Pt(max(base_size - 1, 8))
        label_run.font.bold = True
        label_run.font.color.rgb = CALENDAR_OLIVE

        text_run = note_p.add_run(note["text"])
        text_run.font.name = base_font
        text_run.font.size = Pt(max(base_size - 1, 8))
        text_run.font.bold = bool(note.get("bold"))
        text_run.font.italic = bool(note.get("italic"))

        new_paragraphs.append(note_p)

    # Move new paragraphs from end of doc to anchor position, then remove anchor
    for offset, p in enumerate(new_paragraphs):
        parent.insert(anchor_index + offset, p._p)

    # Remove the anchor (placeholder) paragraph
    parent.remove(p_element)

    return {
        "events": len(events),
        "dates": len(by_date),
        "estimated_height_in": estimated_height,
        "overflow_warning": estimated_height > 7.0,
    }


def generate_bulletin(order: OrderOfWorship) -> Path:
    """Generate a Word bulletin by replacing placeholders in the template."""

    doc = Document(str(TEMPLATE_PATH))
    service_date = date.fromisoformat(order.date)

    # Build the replacement map
    replacements = {
        '{{DATE}}': _ordinal_date(service_date),
        '{{SERVICE_TITLE}}': order.serviceTitle or '',
        '{{OPENING_HYMN_TITLE}}': order.praiseHymn1.title if order.praiseHymn1 else 'TBD',
        '{{OFFERTORY_HYMN_TITLE}}': order.praiseHymn2.title if order.praiseHymn2 else 'TBD',
        '{{OFFERTORY_HYMN}}': order.praiseHymn2.title if order.praiseHymn2 else 'TBD',
        '{{DOX}}': order.doxology.number.lstrip('0') if order.doxology else '95',
        '{{CREED}}': order.creed.number.lstrip('0') if order.creed else '881',
        '{{CREED_TITLE}}': order.creed.title if order.creed else '',
        '{{PRAYER_HYMN_NUMBER}}': order.prayerHymn.number.lstrip('0') if order.prayerHymn else '',
        '{{PRAYER_HYMN_TITLE}}': order.prayerHymn.title if order.prayerHymn else '',
        '{{LITURGICAL_PRAYER}}': _liturgical_prayer_label(order),
        '{{SCRIPTURE}}': order.scripture or '',
        '{{SPEAKER}}': order.speakerShortName or '',
        '{{SERMON_TITLE}}': order.sermonTitle or '',
        '{{SERMON_SUBTITLE}}': order.sermonSubtitle or '',
        '{{CLOSING_HYMN_NUMBER}}': order.closingHymn.number.lstrip('0') if order.closingHymn else '',
        '{{CLOSING_HYMN_TITLE}}': order.closingHymn.title if order.closingHymn else '',
    }

    # Replace placeholders in all paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                _replace_in_runs(paragraph, placeholder, value)

    # Also check tables (QR code section might be in one)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in replacements.items():
                        if placeholder in paragraph.text:
                            _replace_in_runs(paragraph, placeholder, value)

    # Replace theme image
    theme_path = OUTPUT_DIR / order.themeImageFilename if order.themeImageFilename else None
    _replace_theme_image(doc, theme_path)

    # Render calendar block (if placeholder exists in template)
    calendar_anchor = None
    for paragraph in doc.paragraphs:
        if '{{CALENDAR_BLOCK}}' in paragraph.text:
            calendar_anchor = paragraph
            break

    if calendar_anchor:
        _render_calendar_block(doc, calendar_anchor, order.date)

    # Save
    OUTPUT_DIR.mkdir(exist_ok=True)
    filename = f"{order.date} - Bulletin.docx"
    filepath = OUTPUT_DIR / filename
    doc.save(str(filepath))
    return filepath
