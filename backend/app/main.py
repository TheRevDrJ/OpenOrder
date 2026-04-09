"""Order of Worship — FastAPI backend."""

import json
import sys
from datetime import date, timedelta
from pathlib import Path

from fastapi import FastAPI, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles

from .hymnal import search_hymns, get_hymn, get_hymn_by_ref
from .models import OrderOfWorship
from .bulletin import generate_bulletin
from .slides import generate_slides
from .scripture import fetch_scripture, get_available_translations, parse_reference

app = FastAPI(title="Order of Worship")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

from .paths import OUTPUT_DIR, FRONTEND_DIST_DIR, get_settings, update_data_dir, _data_root


def next_sunday() -> str:
    """Return the date of the upcoming Sunday (or today if it's Sunday)."""
    today = date.today()
    days_until_sunday = (6 - today.weekday()) % 7
    if days_until_sunday == 0 and today.weekday() != 6:
        days_until_sunday = 7
    return (today + timedelta(days=days_until_sunday)).isoformat()


# --- Health / Info ---

@app.get("/api/health")
def health():
    return {"status": "ok", "nextSunday": next_sunday()}


# --- Hymnal ---

@app.get("/api/hymnal/search")
def hymnal_search(q: str = "", limit: int = 20):
    return search_hymns(q, limit)


@app.get("/api/hymnal/{source}/{number}")
def hymnal_get(source: str, number: str):
    hymn = get_hymn_by_ref(source, number)
    if not hymn:
        raise HTTPException(404, "Hymn not found")
    return hymn


# --- Services (save/load) ---

def _service_path(service_date: str) -> Path:
    return OUTPUT_DIR / f"{service_date} - Raw.json"


@app.get("/api/services")
def list_services():
    """List all saved services."""
    files = sorted(OUTPUT_DIR.glob("* - Raw.json"), reverse=True)
    services = []
    for f in files:
        date_str = f.name.split(" - Raw.json")[0]
        services.append({"date": date_str, "filename": f.name})
    return services


@app.get("/api/services/{service_date}")
def get_service(service_date: str):
    path = _service_path(service_date)
    if not path.exists():
        raise HTTPException(404, "Service not found")
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


@app.post("/api/services/{service_date}")
def save_service(service_date: str, data: OrderOfWorship):
    data.date = service_date
    path = _service_path(service_date)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data.model_dump(), f, indent=2, ensure_ascii=False)
    return {"saved": True, "path": str(path)}


# --- Theme image upload ---

@app.post("/api/services/{service_date}/theme-image")
async def upload_theme_image(service_date: str, file: UploadFile):
    ext = Path(file.filename).suffix.lower() or ".jpg"
    content = await file.read()

    # Convert unsupported formats (WebP, etc.) to PNG at upload time
    supported = {'.bmp', '.gif', '.jpg', '.jpeg', '.png', '.tiff', '.tif', '.wmf'}
    if ext not in supported:
        from PIL import Image
        import io
        img = Image.open(io.BytesIO(content))
        buf = io.BytesIO()
        img.save(buf, format='PNG')
        content = buf.getvalue()
        ext = '.png'

    filename = f"{service_date} - Theme{ext}"
    dest = OUTPUT_DIR / filename
    with open(dest, "wb") as f:
        f.write(content)
    return {"filename": filename}


# --- Generate ---

@app.post("/api/generate/bulletin/{service_date}")
def gen_bulletin(service_date: str):
    path = _service_path(service_date)
    if not path.exists():
        raise HTTPException(404, "Service not found — save first")
    with open(path, "r", encoding="utf-8") as f:
        data = OrderOfWorship(**json.load(f))
    try:
        filepath = generate_bulletin(data)
    except PermissionError:
        raise HTTPException(
            409, "The bulletin file is open in another program (probably Word). Close it and try again."
        )
    except Exception as e:
        import traceback, logging
        logging.getLogger("OpenOrder").error("Bulletin generation failed:\n%s", traceback.format_exc())
        raise HTTPException(500, f"Bulletin generation failed: {str(e)}")
    return {"filename": filepath.name}


@app.post("/api/generate/slides/{service_date}")
def gen_slides(service_date: str):
    path = _service_path(service_date)
    if not path.exists():
        raise HTTPException(404, "Service not found — save first")
    with open(path, "r", encoding="utf-8") as f:
        data = OrderOfWorship(**json.load(f))
    try:
        filepath = generate_slides(data)
    except PermissionError:
        raise HTTPException(
            409, "The slides file is open in another program (probably PowerPoint). Close it and try again."
        )
    except Exception as e:
        import traceback, logging
        logging.getLogger("OpenOrder").error("Slide generation failed:\n%s", traceback.format_exc())
        raise HTTPException(500, f"Slide generation failed: {str(e)}")
    return {"filename": filepath.name}


# --- Scripture ---

# --- Template management ---

EXPECTED_PLACEHOLDERS = [
    '{{DATE}}', '{{SERVICE_TITLE}}',
    '{{OPENING_HYMN_TITLE}}', '{{OFFERTORY_HYMN_TITLE}}', '{{OFFERTORY_HYMN}}',
    '{{DOX}}', '{{CREED}}', '{{CREED_TITLE}}',
    '{{PRAYER_HYMN_NUMBER}}', '{{PRAYER_HYMN_TITLE}}',
    '{{LITURGICAL_PRAYER}}',
    '{{SCRIPTURE}}', '{{SPEAKER}}',
    '{{SERMON_TITLE}}', '{{SERMON_SUBTITLE}}',
    '{{CLOSING_HYMN_NUMBER}}', '{{CLOSING_HYMN_TITLE}}',
]


@app.get("/api/template/info")
def template_info():
    """Return info about the current bulletin template."""
    from .bulletin import TEMPLATE_PATH
    from docx import Document

    if not TEMPLATE_PATH.exists():
        return {"exists": False, "name": None, "placeholders": []}

    doc = Document(str(TEMPLATE_PATH))
    full_text = '\n'.join(p.text for p in doc.paragraphs)

    found = [p for p in EXPECTED_PLACEHOLDERS if p in full_text]
    missing = [p for p in EXPECTED_PLACEHOLDERS if p not in full_text]

    return {
        "exists": True,
        "name": TEMPLATE_PATH.name,
        "found": found,
        "missing": missing,
        "total_expected": len(EXPECTED_PLACEHOLDERS),
    }


@app.post("/api/template/upload")
async def upload_template(file: UploadFile):
    """Upload and validate a new bulletin template."""
    if not file.filename.endswith('.docx'):
        raise HTTPException(400, "Template must be a .docx file")

    content = await file.read()

    # Validate it's a real docx by trying to open it
    try:
        import io
        from docx import Document
        doc = Document(io.BytesIO(content))
    except Exception:
        raise HTTPException(400, "Invalid .docx file — could not parse")

    # Check for placeholders
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    found = [p for p in EXPECTED_PLACEHOLDERS if p in full_text]
    missing = [p for p in EXPECTED_PLACEHOLDERS if p not in full_text]

    # Save it
    from .bulletin import TEMPLATE_PATH
    TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
    with open(TEMPLATE_PATH, "wb") as f:
        f.write(content)

    return {
        "saved": True,
        "name": file.filename,
        "found": found,
        "missing": missing,
        "total_expected": len(EXPECTED_PLACEHOLDERS),
    }


# --- Scripture ---

@app.get("/api/scripture/translations")
def scripture_translations():
    return get_available_translations()


@app.get("/api/scripture/fetch")
def scripture_fetch(ref: str = "", translation: str = "BSB"):
    if not ref.strip():
        raise HTTPException(400, "Scripture reference is required")
    parsed = parse_reference(ref)
    if not parsed:
        raise HTTPException(400, f"Could not parse scripture reference: {ref}")
    data = fetch_scripture(ref, translation)
    if not data:
        raise HTTPException(404, f"Could not fetch scripture for {ref} ({translation})")
    return data


# --- Settings ---

@app.get("/api/settings")
def api_get_settings():
    settings = get_settings()
    settings["data_dir_current"] = str(_data_root())
    return settings


@app.post("/api/settings/data-dir")
def api_set_data_dir(body: dict):
    new_dir = body.get("data_dir", "").strip()
    if not new_dir:
        raise HTTPException(400, "data_dir is required")
    # Normalize path — fix double backslashes, convert forward slashes
    new_dir = str(Path(new_dir).resolve())
    p = Path(new_dir)
    if not p.exists():
        raise HTTPException(400, f"Directory does not exist: {new_dir}")
    update_data_dir(new_dir)
    # Reload the hymnal index since the data dir changed
    from .hymnal import _load_index
    _load_index()
    return {"data_dir": new_dir, "status": "ok"}


# --- File downloads ---

@app.get("/api/download/{filename}")
def download_file(filename: str):
    path = OUTPUT_DIR / filename
    if not path.exists():
        raise HTTPException(404, "File not found")
    return FileResponse(path, filename=filename)


# --- Serve frontend in production ---
FRONTEND_DIST = FRONTEND_DIST_DIR
if FRONTEND_DIST.exists():
    # Serve static assets (JS, CSS, images)
    app.mount("/assets", StaticFiles(directory=FRONTEND_DIST / "assets"), name="static-assets")
    # Serve other static files (favicon, etc.)
    for static_file in FRONTEND_DIST.iterdir():
        if static_file.is_file() and static_file.name != "index.html":
            @app.get(f"/{static_file.name}")
            def serve_static(f=static_file):
                return FileResponse(f)

    # SPA fallback — serve index.html for any non-API route
    @app.get("/{full_path:path}")
    def serve_spa(full_path: str):
        # Don't intercept API routes
        if full_path.startswith("api/"):
            raise HTTPException(404)
        return FileResponse(FRONTEND_DIST / "index.html")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8316)
